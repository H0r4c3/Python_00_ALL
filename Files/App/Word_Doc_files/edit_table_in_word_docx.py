'https://stackoverflow.com/questions/54890428/edit-table-in-word-docx-using-python'


from docx import Document

path = r'C:\Users\Horace.000\eclipse-workspace\Python_Project_6_Online_Courses\00_ALL\Files\Word_Doc_files\my_output_file.docx'

doc = Document(path)
doc.tables #a list of all tables in document
print("Retrieved value: " + doc.tables[0].cell(0, 0).text)
doc.tables[0].cell(0, 0).text = "new value"
doc.tables[0].add_row() #ADD ROW HERE

doc.save(path)